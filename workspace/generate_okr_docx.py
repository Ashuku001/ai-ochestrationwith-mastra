from docx import Document
from docx.shared import Inches

def create_okr_docx(filename):
    document = Document()

    document.add_heading('OKR Report: Financial Resilience', 0)

    document.add_heading('Objective', level=1)
    p = document.add_paragraph()
    p.add_run('Title: ').bold = True
    p.add_run('Build a highly automated backend ecosystem that drives institutional efficiency and resilience.')
    
    p = document.add_paragraph()
    p.add_run('Owner: ').bold = True
    p.add_run('Ashuku Ezra')
    
    p = document.add_paragraph()
    p.add_run('Cycle: ').bold = True
    p.add_run('1000 orders by end of the year')
    
    p = document.add_paragraph()
    p.add_run('Priority: ').bold = True
    p.add_run('Medium')
    
    p = document.add_paragraph()
    p.add_run('Status: ').bold = True
    p.add_run('Not Started')

    document.add_heading('Key Results', level=1)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Key Result'
    hdr_cells[1].text = 'Target'
    hdr_cells[2].text = 'Status'
    
    krs = [
        ("Automate 80% of manual data reconciliation tasks between the Student Record System and Finance systems.", "80%", "Not Started"),
        ("Reduce cloud infrastructure spend by 10% through the implementation of auto-scaling and spot instance usage.", "10%", "Not Started"),
        ("Resolve 100% of 'Critical' and 'High' security vulnerabilities within 48 hours of detection.", "100%", "Not Started")
    ]
    
    for kr, target, status in krs:
        row_cells = table.add_row().cells
        row_cells[0].text = kr
        row_cells[1].text = target
        row_cells[2].text = status

    document.add_heading('Summary', level=1)
    document.add_paragraph(
        'This objective directly supports the university\'s strategic goal of Financial Resilience by focusing on operational efficiency, cost reduction, and risk mitigation through backend automation.'
    )

    document.save(filename)

if __name__ == "__main__":
    create_okr_docx("Financial_Resilience_OKR.docx")
